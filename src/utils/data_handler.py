"""
Data handling utilities for managing scan history and generating reports.
"""
import json
import pandas as pd
from docx import Document
from io import BytesIO
from src.config import HISTORY_FILE


def save_history(history):
    """
    Save analysis history to JSON file.
    
    Args:
        history (list): List of scan history records
    """
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f, indent=4)


def load_history():
    """
    Load analysis history from JSON file.
    
    Returns:
        list: List of scan history records
    """
    try:
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return []


def generate_report_data(history):
    """
    Generate report data from history.
    
    Args:
        history (list): Scan history records
        
    Returns:
        tuple: (report_data list, status_counts DataFrame)
    """
    report_data = []

    for record in history:
        sender_email = record.get('From', 'N/A')
        date_time = record.get('timestamp', 'N/A')
        subject = record.get('Subject', 'N/A')

        # Process URLs
        for url in record.get('URLs', []):
            url_scan_results = record.get('URL Scan Results', {}).get(url, {})
            malicious_count = url_scan_results.get('malicious', 0)
            suspicious_count = url_scan_results.get('suspicious', 0)

            if malicious_count > 0:
                status = "Malicious"
            elif suspicious_count > 0:
                status = "Suspicious"
            else:
                status = "Harmless"

            report_data.append({
                "Sender": sender_email,
                "DateTime": date_time,
                "Subject": subject,
                "Type": "URL",
                "Value": url,
                "Status": status
            })

        # Process IPs
        for ip in record.get('IPs', []):
            ip_scan_results = record.get('IP Scan Results', {}).get(ip, {})
            malicious_count = ip_scan_results.get('malicious', 0)
            suspicious_count = ip_scan_results.get('suspicious', 0)

            if malicious_count > 0:
                status = "Malicious"
            elif suspicious_count > 0:
                status = "Suspicious"
            else:
                status = "Harmless"

            report_data.append({
                "Sender": sender_email,
                "DateTime": date_time,
                "Subject": subject,
                "Type": "IP",
                "Value": ip,
                "Status": status
            })

    # Convert to DataFrame
    df = pd.DataFrame(report_data)
    
    # Generate status counts
    if not df.empty:
        status_counts = df["Status"].value_counts().reset_index()
        status_counts.columns = ["Status", "Count"]
        
        # Ensure all status types are present
        for status in ["Malicious", "Suspicious", "Harmless"]:
            if status not in status_counts["Status"].values:
                new_row = pd.DataFrame({"Status": [status], "Count": [0]})
                status_counts = pd.concat([status_counts, new_row], ignore_index=True)
    else:
        status_counts = pd.DataFrame({"Status": ["Malicious", "Suspicious", "Harmless"], "Count": [0, 0, 0]})

    return report_data, status_counts, df


def generate_word_report(df, status_counts):
    """
    Generate a Word document report.
    
    Args:
        df (DataFrame): Report data
        status_counts (DataFrame): Status count summary
        
    Returns:
        BytesIO: Word document in memory
    """
    doc = Document()
    doc.add_heading("Email Security Analysis Report", level=1)
    doc.add_paragraph("This report summarizes the analysis of emails, identifying malicious links and IPs.")

    doc.add_heading("Summary", level=2)
    for _, row in status_counts.iterrows():
        doc.add_paragraph(f"{row['Status']}: {row['Count']} occurrences")

    doc.add_heading("Details", level=2)
    for _, row in df.iterrows():
        doc.add_paragraph(f"Sender: {row['Sender']}")
        doc.add_paragraph(f"Date and Time: {row['DateTime']}")
        doc.add_paragraph(f"Subject: {row['Subject']}")
        doc.add_paragraph(f"Type: {row['Type']}")
        doc.add_paragraph(f"Value: {row['Value']}")
        doc.add_paragraph(f"Status: {row['Status']}")
        doc.add_paragraph("------")

    # Save to BytesIO
    report_file = BytesIO()
    doc.save(report_file)
    report_file.seek(0)
    
    return report_file
