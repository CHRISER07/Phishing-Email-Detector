import streamlit as st
import json
import os
from docx import Document
from io import BytesIO
import pandas as pd
import plotly.express as px
import imaplib
import email
from email.header import decode_header
from bs4 import BeautifulSoup
import re
import requests
import time
import concurrent.futures
import yagmail
from dotenv import load_dotenv


load_dotenv()

# VirusTotal API Key
VIRUSTOTAL_API_KEY = '6521a614ec622362884db41bf42559c8826e21cfe12f9f97450f992e79000824'

# Email Server Configuration
IMAP_SERVERS = {
    "Gmail": "imap.gmail.com",
    "Outlook": "imap-mail.outlook.com",
    "Yahoo": "imap.mail.yahoo.com"
}

# Directory for storing scanned results
RESULTS_DIR = "scan_results"
os.makedirs(RESULTS_DIR, exist_ok=True)

# File to store history
HISTORY_FILE = "history.json"

def save_history(history):
    """Save analysis history to JSON file"""
    with open(HISTORY_FILE, "w") as f:
        json.dump(history, f, indent=4)

def load_history():
    """Load analysis history from JSON file"""
    try:
        with open(HISTORY_FILE, "r") as f:
            return json.load(f)
    except FileNotFoundError:
        return []

def fetch_emails(email_user, app_password, provider, folder="INBOX", num_emails=5):
    try:
        imap_server = IMAP_SERVERS.get(provider)
        if not imap_server:
            return f"Unsupported email provider: {provider}"

        mail = imaplib.IMAP4_SSL(imap_server, 993)
        mail.login(email_user, app_password)
        mail.select(folder)

        status, messages = mail.search(None, "ALL")
        mail_ids = messages[0].split()[-num_emails:]

        email_list = []
        for mail_id in reversed(mail_ids):
            status, msg_data = mail.fetch(mail_id, "(RFC822)")
            for response_part in msg_data:
                if isinstance(response_part, tuple):
                    msg = email.message_from_bytes(response_part[1])
                    subject, encoding = decode_header(msg["Subject"])[0]
                    if isinstance(subject, bytes):
                        subject = subject.decode(encoding or "utf-8")
                    sender = msg.get("From")
                    body = extract_email_body(msg)
                    urls = extract_urls(body)
                    ips = extract_ips(body)
                    
                    email_list.append({
                        "Subject": subject,
                        "From": sender,
                        "Body": body,
                        "URLs": urls,
                        "IPs": ips
                    })
        mail.logout()
        return email_list
    except Exception as e:
        return str(e)

def extract_email_body(msg):
    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            content_type = part.get_content_type()
            if content_type == "text/plain":
                body = part.get_payload(decode=True).decode("utf-8", errors="ignore")
                break
            elif content_type == "text/html":
                html = part.get_payload(decode=True).decode("utf-8", errors="ignore")
                soup = BeautifulSoup(html, "html.parser")
                body = soup.get_text()
    else:
        body = msg.get_payload(decode=True).decode("utf-8", errors="ignore")
    return body.strip()

def extract_urls(text):
    """Extract URLs from text using regex"""
    return re.findall(r'https?://\S+', text)

def extract_ips(text):
    """Extract IP addresses from text using regex"""
    return re.findall(r'\b(?:\d{1,3}\.){3}\d{1,3}\b', text)

def check_url_virustotal(url):
    """Check URL against VirusTotal"""
    try:
        response = requests.post(
            "https://www.virustotal.com/api/v3/urls",
            headers={"x-apikey": VIRUSTOTAL_API_KEY},
            data={"url": url}
        )
        if response.status_code != 200:
            return {}
        analysis_id = response.json().get("data", {}).get("id", "")
        if not analysis_id:
            return {}
        
        for _ in range(5):
            time.sleep(2)
            report_response = requests.get(
                f"https://www.virustotal.com/api/v3/analyses/{analysis_id}",
                headers={"x-apikey": VIRUSTOTAL_API_KEY}
            )
            if report_response.status_code == 200:
                break
        
        report = report_response.json()
        stats = report.get("data", {}).get("attributes", {}).get("stats", {})
        return {
            "malicious": stats.get("malicious", 0),
            "suspicious": stats.get("suspicious", 0),
            "harmless": stats.get("harmless", 0)
        }
    except Exception as e:
        return {}

def check_ip_virustotal(ip):
    """Check IP address against VirusTotal"""
    try:
        response = requests.get(
            f"https://www.virustotal.com/api/v3/ip_addresses/{ip}",
            headers={"x-apikey": VIRUSTOTAL_API_KEY}
        )
        if response.status_code != 200:
            return {}
        
        data = response.json().get("data", {})
        stats = data.get("attributes", {}).get("last_analysis_stats", {})
        return {
            "malicious": stats.get("malicious", 0),
            "suspicious": stats.get("suspicious", 0),
            "harmless": stats.get("harmless", 0)
        }
    except Exception as e:
        return {}
    
def malicious_history_page():
    st.header("📜 Malicious Links & IPs History")

    history = load_history()
    if not history:
        st.info("No malicious history found.")
        return

    # Create a list to store table data
    table_data = []

    # Iterate through history and extract relevant data
    for record in history:
        subject = record.get("Subject", "N/A")
        sender = record.get("From", "N/A")
        timestamp = record.get("timestamp", "N/A")
        malicious_urls = record.get("URLs", [])
        malicious_ips = record.get("IPs", [])

        # Add URLs to table data
        for url in malicious_urls:
            table_data.append({
                "Type": "URL",
                "Value": url,
                "Subject": subject,
                "Sender": sender,
                "Timestamp": timestamp,
                "Status": "Malicious"
            })

        # Add IPs to table data
        for ip in malicious_ips:
            table_data.append({
                "Type": "IP",
                "Value": ip,
                "Subject": subject,
                "Sender": sender,
                "Timestamp": timestamp,
                "Status": "Malicious"
            })

    # Convert to DataFrame for better display
    if table_data:
        df = pd.DataFrame(table_data)

        # Reorder columns for better readability
        df = df[["Timestamp", "Sender", "Subject", "Type", "Value", "Status"]]

        # Display the table with Streamlit
        st.dataframe(
            df,
            column_config={
                "Timestamp": "Timestamp",
                "Sender": "Sender",
                "Subject": "Subject",
                "Type": "Type",
                "Value": "Value",
                "Status": "Status"
            },
            use_container_width=True
        )

        # Add a download button for the table
        csv = df.to_csv(index=False).encode('utf-8')
        st.download_button(
            label="Download Malicious History as CSV",
            data=csv,
            file_name="malicious_history.csv",
            mime="text/csv"
        )
    else:
        st.info("No malicious URLs or IPs detected in history.")

def send_security_alert(recipient, subject, body):
    """Send security alert email using app password"""
    try:
        # Initialize yagmail SMTP
        yag = yagmail.SMTP(
            user='jerjacmat@gmail.com',  # Sender's email
            password='cehq ddgp mwtu qybe'  # Sender's app password
        )

        # Send email
        yag.send(
            to=recipient,
            subject=subject,
            contents=body
        )
        return True
    except Exception as e:
        print(f"Error sending email: {e}")
        return False

def main():
    st.title("🔒 Email Security Analyzer")

    # Sidebar for navigation
    page = st.sidebar.selectbox("Navigate",
                                ["Analyze Emails", "Malicious History", "Detailed Reports", "Dashboard"]
                                )

    if page == "Analyze Emails":
        analyze_emails_page()
    elif page == "Malicious History":
        malicious_history_page()
    elif page == "Detailed Reports":
        detailed_reports_page()
    elif page == "Dashboard":
        dashboard_page()

def analyze_emails_page():
    st.header("Email Security Scanning")

    # Email input fields
    email_user = st.text_input("📩 Enter Your Email")
    app_password = st.text_input("🔑 Enter App Password", type="password")
    provider = st.selectbox("🌐 Select Email Provider", ["Gmail", "Outlook", "Yahoo"])
    num_emails = st.slider("📌 Number of Emails to Fetch", 1, 10, 5)

    if st.button("Fetch & Deep Scan Emails"):
        if email_user and app_password and provider:
            with st.spinner("Fetching emails..."):
                emails = fetch_emails(email_user, app_password, provider, num_emails=num_emails)
                if isinstance(emails, list):
                    email_data_list = []
                    with concurrent.futures.ThreadPoolExecutor() as executor:
                        # Check URLs
                        url_futures = {executor.submit(check_url_virustotal, url): url for email_data in emails for url in email_data['URLs']}
                        url_scan_results = {url: future.result() for future, url in zip(concurrent.futures.as_completed(url_futures), url_futures.values())}

                        # Check IPs
                        ip_futures = {executor.submit(check_ip_virustotal, ip): ip for email_data in emails for ip in email_data['IPs']}
                        ip_scan_results = {ip: future.result() for future, ip in zip(concurrent.futures.as_completed(ip_futures), ip_futures.values())}

                    for email_data in emails:
                        malicious_urls = [url for url in email_data['URLs'] if url_scan_results.get(url, {}).get('malicious', 0) > 0]
                        malicious_ips = [ip for ip in email_data['IPs'] if ip_scan_results.get(ip, {}).get('malicious', 0) > 0]

                        if malicious_urls or malicious_ips:
                            history = load_history()
                            history.append({
                                "Subject": email_data["Subject"],
                                "From": email_data["From"],
                                "Body": email_data["Body"],
                                "URLs": malicious_urls,
                                "IPs": malicious_ips,
                                "URL Scan Results": {url: url_scan_results[url] for url in malicious_urls},
                                "IP Scan Results": {ip: ip_scan_results[ip] for ip in malicious_ips},
                                "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
                            })
                            save_history(history)

                            # Send alert button
                            if st.button(f"🚨 Send Alert for: {email_data['Subject']}", key=f"alert_{email_data['Subject']}"):
                                body = f"""Malicious email detected!
                                
                                Subject: {email_data['Subject']}
                                From: {email_data['From']}
                                Detected URLs: {', '.join(malicious_urls)}
                                Detected IPs: {', '.join(malicious_ips)}
                                """
                                
                                if send_security_alert(email_user, "Security Alert!", body):
                                    st.success("Alert sent successfully!")
                                else:
                                    st.error("Failed to send alert")

                    # Display results
                    results_df = pd.DataFrame([{
                        "DateTime": time.strftime("%Y-%m-%d %H:%M:%S"),
                        "Subject": e["Subject"],
                        "From": e["From"],
                        "Body": e["Body"],
                        "Total URLs": len(e["URLs"]),
                        "Malicious URLs": len([u for u in e["URLs"] if url_scan_results.get(u, {}).get('malicious', 0) > 0]),
                        "Total IPs": len(e["IPs"]),
                        "Malicious IPs": len([ip for ip in e["IPs"] if ip_scan_results.get(ip, {}).get('malicious', 0) > 0])
                    } for e in emails])

                    # Display the DataFrame without any additional styling
                    st.dataframe(results_df)
                    
                else:
                    st.error(f"Error: {emails}")
        else:
            st.warning("Please fill all required fields")

# def malicious_history_page():
#     st.header("Malicious Links & IPs History")

#     history = load_history()
#     if not history:
#         st.info("No malicious history found.")
#         return

#     for i, record in enumerate(history, 1):
#         st.subheader(f"Suspicious Email {i}")

#         st.markdown("#### Suspicious Links")
#         for url in record.get('URLs', []):
#             st.write(f"- {url}: **Malicious**")

#         st.markdown("#### Suspicious IPs")
#         for ip in record.get('IPs', []):
#             st.write(f"- {ip}: **Malicious**")

#         st.markdown("#### Scan Results")
#         st.json({
#             "URLs": record.get('URL Scan Results', {}),
#             "IPs": record.get('IP Scan Results', {})
#         })

#         st.divider()

def detailed_reports_page():
    st.header("Detailed Email Analysis Reports")

    history = load_history()
    if not history:
        st.info("No detailed reports available.")
        return

    selected_report = st.selectbox(
        "Select a Report",
        [f"Report {i + 1}" for i in range(len(history))]
    )

    report_index = int(selected_report.split()[-1]) - 1
    selected_history = history[report_index]

    st.subheader("Links Detailed Analysis")
    for url, result in selected_history.get('URL Scan Results', {}).items():
        st.write(f"URL: {url}")
        if result.get('malicious', 0) > 0:
            st.error("🚨 Malicious URL Detected")
            st.json(result)
        else:
            st.success("Safe URL")

    st.subheader("IPs Detailed Analysis")
    for ip, result in selected_history.get('IP Scan Results', {}).items():
        st.write(f"IP: {ip}")
        if result.get('malicious', 0) > 0:
            st.error("🚨 Malicious IP Detected")
            st.json(result)
        else:
            st.success("Safe IP")

def dashboard_page():
    st.header("📊 Dashboard")

    history = load_history()
    if not history:
        st.info("No data available for the dashboard.")
        return

    # Convert history to DataFrame
    df = pd.DataFrame(history)
    
    # Ensure the 'timestamp' column is of datetime type
    df['timestamp'] = pd.to_datetime(df['timestamp'], errors='coerce')

    # Key Metrics
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Total Threats Detected", len(df))
    with col2:
        st.metric("Unique Malicious URLs", df['URLs'].explode().nunique())
    with col3:
        st.metric("Unique Malicious IPs", df['IPs'].explode().nunique())

    # Time-based Analysis
    st.subheader("Detection Timeline")
    time_df = df.set_index('timestamp').resample('D').size().reset_index(name='count')
    fig = px.line(time_df, x='timestamp', y='count', 
                 title="Malicious Activity Over Time",
                 labels={'timestamp': 'Date', 'count': 'Detections'})
    st.plotly_chart(fig)

    # Threat Breakdown
    st.subheader("Threat Composition")
    try:
        # Combine URLs and IPs into a single DataFrame
        url_counts = df['URLs'].explode().value_counts().reset_index()
        url_counts.columns = ['Threat', 'Count']
        url_counts['Type'] = 'URL'

        ip_counts = df['IPs'].explode().value_counts().reset_index()
        ip_counts.columns = ['Threat', 'Count']
        ip_counts['Type'] = 'IP'

        threat_types = pd.concat([url_counts, ip_counts])
        
        col1, col2 = st.columns(2)
        with col1:
            fig = px.pie(threat_types.head(10), names='Threat', values='Count',
                        title="Top Malicious Threats")
            st.plotly_chart(fig)
        
        with col2:
            fig = px.bar(threat_types.head(10), x='Threat', y='Count',
                        title="Most Frequent Threats")
            st.plotly_chart(fig)
    except Exception as e:
        st.error(f"Error generating threat breakdown: {e}")

    # Source Analysis
    st.subheader("Threat Origins")
    try:
        source_df = df['From'].value_counts().reset_index()
        source_df.columns = ['Sender', 'Count']
        fig = px.treemap(source_df, path=['Sender'], values='Count',
                        title="Threats by Sender")
        st.plotly_chart(fig)
    except KeyError:
        st.warning("No sender information available")

    # Add a report generator
    st.subheader("Generate Report")
    if st.button("Download Report"):
        generate_report(history)
        st.success("Report generated successfully!")

def generate_report(history):
    """Generate a downloadable report from the history data"""
    report_data = []

    for record in history:
        sender_email = record.get('From', 'N/A')
        date_time = record.get('timestamp', 'N/A')
        subject = record.get('Subject', 'N/A')

        # Ensure that the DateTime and Sender fields are captured
        if not date_time:
            date_time = "N/A"  # Default to "N/A" if no date is present
        if not sender_email:
            sender_email = "N/A"  # Default to "N/A" if no sender is present
        
        # Process URLs
        for url in record.get('URLs', []):
            url_scan_results = record.get('URL Scan Results', {}).get(url, {})
            malicious_count = url_scan_results.get('malicious', 0)
            suspicious_count = url_scan_results.get('suspicious', 0)
            harmless_count = url_scan_results.get('harmless', 0)

            # Determine status based on the scan results
            if malicious_count > 0:
                status = "Malicious"
            elif suspicious_count > 0:
                status = "Suspicious"
            else:
                status = "Harmless"

            # Add to report data
            report_data.append({
                "Sender": sender_email,
                "DateTime": date_time,
                "Subject": subject,
                "Type": "URL",
                "Value": url,
                "Status": status
            })

        # Process IPs
        for ip in record.get('IPs', []):
            ip_scan_results = record.get('IP Scan Results', {}).get(ip, {})
            malicious_count = ip_scan_results.get('malicious', 0)
            suspicious_count = ip_scan_results.get('suspicious', 0)
            harmless_count = ip_scan_results.get('harmless', 0)

            # Determine status based on the scan results
            if malicious_count > 0:
                status = "Malicious"
            elif suspicious_count > 0:
                status = "Suspicious"
            else:
                status = "Harmless"

            # Add to report data
            report_data.append({
                "Sender": sender_email,
                "DateTime": date_time,
                "Subject": subject,
                "Type": "IP",
                "Value": ip,
                "Status": status
            })

    # Convert to DataFrame for easier handling
    df = pd.DataFrame(report_data)

    # Generate bar chart: counts of each URL/IP status (Malicious, Suspicious, Harmless)
    status_counts = df["Status"].value_counts().reset_index()
    status_counts.columns = ["Status", "Count"]

    # Ensure all status types are present, even if the count is zero
    # Ensure all status types are present, even if the count is zero
    for status in ["Malicious", "Suspicious", "Harmless"]:
        if status not in status_counts["Status"].values:
            # Create a DataFrame with a new row for the missing status
            new_row = pd.DataFrame({"Status": [status], "Count": [0]})
            status_counts = pd.concat([status_counts, new_row], ignore_index=True)


    bar_chart = px.bar(
        status_counts,
        x="Status",
        y="Count",
        color="Status",
        title="URL/IP Status Distribution (Malicious, Suspicious, Harmless)",
        labels={"Count": "Count", "Status": "URL/IP Status"}
    )
    st.plotly_chart(bar_chart)

    # Generate pie chart: proportion of each URL/IP status (Malicious, Suspicious, Harmless)
    pie_chart = px.pie(
        status_counts,
        names="Status",
        values="Count",
        title="Proportion of URL/IP Statuses (Malicious, Suspicious, Harmless)",
        color="Status"
    )
    st.plotly_chart(pie_chart)

    # Save the report to a CSV file
    report_file_csv = "email_security_report.csv"
    df.to_csv(report_file_csv, index=False)

    # Generate a Word document summary
    doc = Document()
    doc.add_heading("Email Security Analysis Report", level=1)

    doc.add_paragraph("This report summarizes the analysis of emails, identifying malicious links and IPs.")

    doc.add_heading("Summary", level=2)
    for index, row in status_counts.iterrows():
        doc.add_paragraph(f"{row['Status']}: {row['Count']} occurrences")

    doc.add_heading("Details", level=2)
    for _, row in df.iterrows():
        doc.add_paragraph(f"Sender: {row['Sender']}")
        doc.add_paragraph(f"Date and Time: {row['DateTime']}")
        doc.add_paragraph(f"Subject: {row['Subject']}")
        doc.add_paragraph(f"Type: {row['Type']}")
        doc.add_paragraph(f"Value: {row['Value']}")
        doc.add_paragraph(f"Status: {row['Status']}")
        doc.add_paragraph("------")

    # Save the Word report in memory
    report_file_doc = BytesIO()
    doc.save(report_file_doc)
    report_file_doc.seek(0)

    # Provide download links
    st.download_button(
        label="Download CSV Report",
        data=open(report_file_csv, "rb").read(),
        file_name=report_file_csv,
        mime="text/csv"
    )

    st.download_button(
        label="Download Word Report",
        data=report_file_doc,
        file_name="email_security_report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )


if __name__ == "__main__":
    main()



